"""
Document text extraction service for multiple file formats.
Supports PDF, DOCX, EPUB, TXT, and Markdown files.
"""
import logging
import os
from io import BytesIO
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extracts text from various document formats."""

    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.epub': 'application/epub+zip',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
    }

    def extract(self, file_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from a file.

        Args:
            file_bytes: Raw file content as bytes
            filename: Original filename (used to determine format)

        Returns:
            Tuple of (extracted_text, error_message)
            If successful, error_message is None.
        """
        ext = self._get_extension(filename).lower()

        if ext not in self.SUPPORTED_FORMATS:
            supported = ', '.join(self.SUPPORTED_FORMATS.keys())
            return "", f"Unsupported file format: {ext}. Supported: {supported}"

        try:
            if ext == '.pdf':
                return self._extract_pdf(file_bytes), None
            elif ext == '.docx':
                return self._extract_docx(file_bytes), None
            elif ext == '.epub':
                return self._extract_epub(file_bytes), None
            elif ext in ('.txt', '.md'):
                return self._extract_text(file_bytes), None
        except Exception as e:
            logger.error(f"Error extracting {filename}: {e}")
            return "", f"Failed to extract text: {str(e)}"

        return "", "Unknown error during extraction"

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pypdf."""
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(file_bytes))
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Error extracting page {page_num}: {e}")
                continue

        return "\n\n".join(text_parts)

    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document

        doc = Document(BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)

    def _extract_epub(self, file_bytes: bytes) -> str:
        """Extract text from EPUB using ebooklib."""
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(BytesIO(file_bytes))
        text_parts = []

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                try:
                    soup = BeautifulSoup(item.get_content(), 'lxml')

                    # Remove script and style elements
                    for element in soup(['script', 'style']):
                        element.decompose()

                    text = soup.get_text(separator='\n')
                    # Clean up whitespace
                    lines = [line.strip() for line in text.split('\n')]
                    text = '\n'.join(line for line in lines if line)

                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting EPUB item: {e}")
                    continue

        return "\n\n".join(text_parts)

    def _extract_text(self, file_bytes: bytes) -> str:
        """Extract text from plain text or markdown files."""
        # Try multiple encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        # Last resort: decode with errors replaced
        return file_bytes.decode('utf-8', errors='replace')

    @staticmethod
    def _get_extension(filename: str) -> str:
        """Get file extension from filename."""
        return os.path.splitext(filename)[1]

    @classmethod
    def get_supported_formats(cls) -> list:
        """Return list of supported file extensions."""
        return list(cls.SUPPORTED_FORMATS.keys())
